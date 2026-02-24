#!/usr/bin/env python3
"""
STAGE 0: Universal AI-powered file parser

Supported formats:
  Tabular:      .xlsx, .csv, .tsv  → AI maps columns
  Documents:    .pdf, .docx, .txt, .md, .rtf  → AI extracts data from text
  (Future:      images → vision API)

Two modes:
  parse_key_file(path, api_key)      → (DataFrame, mapping, warnings, errors)
  parse_answers_file(path, api_key)  → (DataFrame, mapping, warnings, errors)

Usage:
  python3 stage0.py key raw_key.pdf          → Key.xlsx
  python3 stage0.py answers raw_answers.xlsx → Answers.xlsx
"""

import json
import os
import sys
import re
import io
import pandas as pd

# ═══════════════════════════════════════════════════════════════
# CONFIG
# ═══════════════════════════════════════════════════════════════

DEFAULT_MODEL = "claude-sonnet-4-5-20250929"
MAX_SAMPLE_ROWS = 3
MAX_HEADER_CHARS = 200
MAX_TEXT_CHARS = 50000  # max text to send to AI for unstructured parsing

TABULAR_EXTENSIONS = {".xlsx", ".xls", ".csv", ".tsv"}
DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".rtf"}

# ═══════════════════════════════════════════════════════════════
# TEXT EXTRACTION (any format → raw text)
# ═══════════════════════════════════════════════════════════════

def extract_text_from_pdf(filepath):
    """Extract text from PDF using pdfplumber. Detects image-based PDFs."""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("Install pdfplumber: pip3 install pdfplumber")

    pages = []
    total_images = 0

    with pdfplumber.open(filepath) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            # Also try extracting tables
            tables = page.extract_tables() or []
            table_text = ""
            for table in tables:
                for row in table:
                    cells = [str(c or "") for c in row]
                    table_text += " | ".join(cells) + "\n"
            combined = text
            if table_text and table_text not in text:
                combined += "\n[TABLE]\n" + table_text

            # Count images
            if hasattr(page, 'images'):
                total_images += len(page.images)

            pages.append(combined)

    full_text = "\n\n--- PAGE BREAK ---\n\n".join(pages)

    # Check if PDF is image-based (no text but has images or pages)
    text_chars = sum(len(p.strip()) for p in pages)
    if text_chars < 50 and len(pages) > 0:
        raise ValueError(
            f"This PDF appears to be image-based (scanned). "
            f"Found {len(pages)} pages but only {text_chars} characters of text. "
            f"Please convert it to text first (e.g. using OCR), "
            f"or copy-paste the content into a .txt file and upload that instead."
        )

    return full_text


def extract_text_from_docx(filepath):
    """Extract text from Word document."""
    from docx import Document
    doc = Document(filepath)
    parts = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables
    for table in doc.tables:
        parts.append("[TABLE]")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_text_from_txt(filepath):
    """Read plain text file."""
    encodings = ["utf-8", "utf-8-sig", "cp1251", "latin-1", "iso-8859-1"]
    for enc in encodings:
        try:
            with open(filepath, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise ValueError(f"Cannot decode text file: {filepath}")


def extract_text(filepath):
    """Extract text from any supported format. Returns (text, format_type)."""
    ext = os.path.splitext(filepath)[1].lower()

    if ext in (".pdf",):
        return extract_text_from_pdf(filepath), "pdf"
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(filepath), "docx"
    elif ext in (".txt", ".md", ".rtf"):
        return extract_text_from_txt(filepath), "text"
    else:
        raise ValueError(f"Unsupported format: {ext}")


def get_file_type(filepath):
    """Return 'tabular' or 'document' based on extension."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in TABULAR_EXTENSIONS:
        return "tabular"
    elif ext in DOCUMENT_EXTENSIONS:
        return "document"
    else:
        raise ValueError(
            f"Unsupported file format: {ext}\n"
            f"Supported: {', '.join(sorted(TABULAR_EXTENSIONS | DOCUMENT_EXTENSIONS))}"
        )


# ═══════════════════════════════════════════════════════════════
# TABULAR PATH: Read spreadsheet → AI maps columns
# ═══════════════════════════════════════════════════════════════

def read_tabular(filepath):
    """Read xlsx/csv into DataFrame. Returns (df, metadata)."""
    ext = os.path.splitext(filepath)[1].lower()

    if ext in (".csv", ".tsv"):
        sep = "\t" if ext == ".tsv" else None
        for delimiter in ([sep] if sep else [",", ";", "\t"]):
            try:
                df = pd.read_csv(filepath, sep=delimiter, dtype=str, na_filter=False)
                if len(df.columns) > 1:
                    break
            except Exception:
                continue
        else:
            raise ValueError(f"Cannot parse CSV: {filepath}")
    else:
        df = pd.read_excel(filepath, dtype=str, na_filter=False)

    metadata = {
        "filename": os.path.basename(filepath),
        "rows": len(df), "cols": len(df.columns),
        "columns": df.columns.tolist(),
    }
    return df, metadata


def extract_structure(df, max_rows=MAX_SAMPLE_ROWS):
    """Extract headers and sample rows for AI analysis."""
    headers = [str(c).strip()[:MAX_HEADER_CHARS] for c in df.columns]
    samples = []
    for i in range(min(max_rows, len(df))):
        row = {}
        for col, header in zip(df.columns, headers):
            val = str(df.iloc[i][col]).strip()[:200]
            row[header] = val
        samples.append(row)
    return headers, samples


# ── Tabular prompts ──

def _tabular_answers_prompt(headers, samples, metadata):
    sample_text = json.dumps(samples, ensure_ascii=False, indent=2)
    return f"""Analyze this spreadsheet structure and create a column mapping.

TARGET FORMAT: Answers.xlsx with columns:
  Name | Surname | Language | Q1 | Q2 | ... | QN

SOURCE FILE: "{metadata['filename']}" ({metadata['rows']} rows, {metadata['cols']} columns)

COLUMN HEADERS:
{json.dumps(headers, ensure_ascii=False)}

SAMPLE ROWS (first {len(samples)}):
{sample_text}

TASK: Map source columns to target fields. Return ONLY a JSON object:

{{
  "name_col": "<header of column containing first name, or null>",
  "surname_col": "<header of column containing surname/last name, or null>",
  "fullname_col": "<header if name+surname are in ONE column, or null>",
  "language_col": "<header of column containing language, or null>",
  "question_mapping": {{
    "1": "<header of column containing answer to Q1>",
    "2": "<header of column containing answer to Q2>"
  }},
  "skip_cols": ["<headers of columns to ignore, e.g. Timestamp, Email, Score>"],
  "detected_language": "<if no language column, what language are the answers in? or 'mixed'>",
  "notes": "<any issues or ambiguities>"
}}

RULES:
- Question numbers should be inferred from column order or from numbers in headers
- If headers contain full question text (Google Forms), extract the question number from position
- If no separate Name/Surname columns exist but there's a full name column, use fullname_col
- Skip metadata columns (timestamps, emails, scores, IDs)
- Return ONLY valid JSON, no other text"""


def _tabular_key_prompt(headers, samples, metadata):
    sample_text = json.dumps(samples, ensure_ascii=False, indent=2)
    return f"""Analyze this spreadsheet and create a column mapping for an answer key.

TARGET FORMAT: Key.xlsx with columns:
  Number | Question | Answer | Points | Comment

SOURCE FILE: "{metadata['filename']}" ({metadata['rows']} rows, {metadata['cols']} columns)

COLUMN HEADERS:
{json.dumps(headers, ensure_ascii=False)}

SAMPLE ROWS (first {len(samples)}):
{sample_text}

TASK: Map source columns to target fields. Return ONLY a JSON object:

{{
  "number_col": "<header of column with question numbers, or null if sequential>",
  "question_col": "<header of column with question text>",
  "answer_col": "<header of column with correct answers>",
  "points_col": "<header of column with point values, or null>",
  "comment_col": "<header of column with comments/notes, or null>",
  "default_points": <default point value if no points column, e.g. 6>,
  "notes": "<any issues or ambiguities>"
}}

Return ONLY valid JSON, no other text"""


# ── Tabular transformations ──

def apply_answers_mapping(df, mapping):
    result = pd.DataFrame()

    if mapping.get("fullname_col"):
        full_col = mapping["fullname_col"]
        names = df[full_col].str.strip().str.split(r'\s+', n=1, expand=True)
        result["Name"] = names[0] if 0 in names.columns else ""
        result["Surname"] = names[1] if 1 in names.columns else ""
    else:
        nc = mapping.get("name_col")
        sc = mapping.get("surname_col")
        result["Name"] = df[nc].str.strip() if nc and nc in df.columns else ""
        result["Surname"] = df[sc].str.strip() if sc and sc in df.columns else ""

    lc = mapping.get("language_col")
    if lc and lc in df.columns:
        result["Language"] = df[lc].str.strip()
    else:
        result["Language"] = mapping.get("detected_language", "Unknown")

    for qn in sorted(mapping.get("question_mapping", {}).keys(), key=int):
        src = mapping["question_mapping"][qn]
        if src in df.columns:
            result[f"Q{qn}"] = df[src].fillna("").astype(str).str.strip()
        else:
            matched = next((c for c in df.columns if src[:50] in str(c)[:50]), None)
            if matched:
                result[f"Q{qn}"] = df[matched].fillna("").astype(str).str.strip()
            else:
                result[f"Q{qn}"] = ""

    for col in result.columns:
        result[col] = result[col].replace({"None": "", "nan": "", "NaN": "", "-": "", "—": ""})
    result = result[result["Name"].str.strip() != ""].reset_index(drop=True)
    return result


def apply_key_mapping(df, mapping):
    result = pd.DataFrame()

    nc = mapping.get("number_col")
    if nc and nc in df.columns:
        result["Number"] = pd.to_numeric(df[nc], errors="coerce").astype("Int64")
    else:
        result["Number"] = range(1, len(df) + 1)

    qc = mapping.get("question_col")
    result["Question"] = df[qc].str.strip() if qc and qc in df.columns else ""

    ac = mapping.get("answer_col")
    result["Answer"] = df[ac].str.strip() if ac and ac in df.columns else ""

    pc = mapping.get("points_col")
    if pc and pc in df.columns:
        result["Points"] = pd.to_numeric(df[pc], errors="coerce").fillna(
            mapping.get("default_points", 6)).astype(int)
    else:
        result["Points"] = mapping.get("default_points", 6)

    cc = mapping.get("comment_col")
    result["Comment"] = df[cc].fillna("").astype(str).str.strip() if cc and cc in df.columns else ""

    result = result[result["Question"].str.strip() != ""].reset_index(drop=True)
    return result


# ═══════════════════════════════════════════════════════════════
# DOCUMENT PATH: Extract text → AI returns structured data
# ═══════════════════════════════════════════════════════════════

def _truncate_text(text, max_chars=MAX_TEXT_CHARS):
    if len(text) <= max_chars:
        return text
    half = max_chars // 2
    return text[:half] + f"\n\n... [{len(text) - max_chars} characters truncated] ...\n\n" + text[-half:]


def _document_key_prompt(text, filename):
    text = _truncate_text(text)
    return f"""Extract an answer key from the following document.

SOURCE: "{filename}"

DOCUMENT TEXT:
---
{text}
---

TASK: Extract ALL questions with their correct answers from this document.
Return ONLY a JSON object with this structure:

{{
  "questions": [
    {{
      "number": 1,
      "question": "Full question text",
      "answer": "Correct answer",
      "points": 6,
      "comment": "Any special notes, alternative answers, grading rules"
    }},
    ...
  ],
  "notes": "Any issues or observations about the document"
}}

RULES:
- Number questions sequentially if no numbers in the document
- For True/False questions, answer should be "T" or "F"
- For multiple choice, answer should be the letter (A, B, C, D)
- For open questions, answer should be the expected text
- Default points to 6 if not specified in the document
- Put any grading hints in the "comment" field (e.g. "Alt: Moses" for alternative answers)
- Return ONLY valid JSON, no other text"""


def _document_answers_prompt(text, filename):
    text = _truncate_text(text)
    return f"""Extract participant answers from the following document.

SOURCE: "{filename}"

DOCUMENT TEXT:
---
{text}
---

TASK: Extract ALL participant answers from this document.
Return ONLY a JSON object with this structure:

{{
  "participants": [
    {{
      "name": "First name",
      "surname": "Last name",
      "language": "Language of their answers (e.g. Deutsch, English, Русский, Italiano)",
      "answers": {{
        "1": "Their answer to Q1",
        "2": "Their answer to Q2",
        ...
      }}
    }},
    ...
  ],
  "total_questions": 50,
  "notes": "Any issues or observations about the document"
}}

RULES:
- Extract every participant you can find
- Empty or missing answers should be ""
- Detect the language from the answers themselves if not stated
- If names are in "Last, First" format, split them correctly
- If only full names exist, put the first word as name, rest as surname
- CRITICAL — answer format must be minimal to save space:
  * True/False questions: ONLY "T" or "F" (not "True", "False", "Richtig", "Falsch", etc.)
  * Multiple choice questions: ONLY the letter "A", "B", "C", or "D" (strip the dot, text, parentheses)
  * Open questions: the participant's answer text as-is, but trimmed
- Return ONLY valid JSON, no other text"""


def _json_from_document(text, filename, file_type, api_key):
    """Send document text to AI, get structured JSON back."""
    if file_type == "key":
        prompt = _document_key_prompt(text, filename)
    else:
        prompt = _document_answers_prompt(text, filename)

    response = call_api(prompt, api_key)
    return _parse_json_response(response)


def document_to_key_df(data):
    """Convert AI-extracted key JSON to clean DataFrame."""
    questions = data.get("questions", [])
    if not questions:
        return pd.DataFrame(), [], ["No questions extracted from document"]

    rows = []
    for q in questions:
        rows.append({
            "Number": int(q.get("number", 0)),
            "Question": str(q.get("question", "")),
            "Answer": str(q.get("answer", "")),
            "Points": int(q.get("points", 6)),
            "Comment": str(q.get("comment", "")),
        })

    df = pd.DataFrame(rows)
    warnings = []
    errors = []

    if len(df) == 0:
        errors.append("No questions extracted")
    empty_answers = (df["Answer"].str.strip() == "").sum()
    if empty_answers:
        warnings.append(f"{empty_answers} questions with empty answers")

    notes = data.get("notes", "")
    if notes:
        warnings.append(f"AI notes: {notes}")

    return df, warnings, errors


def document_to_answers_df(data):
    """Convert AI-extracted answers JSON to clean DataFrame."""
    participants = data.get("participants", [])
    if not participants:
        return pd.DataFrame(), [], ["No participants extracted from document"]

    total_q = data.get("total_questions", 0)
    # Determine max question number
    all_q_nums = set()
    for p in participants:
        all_q_nums.update(int(k) for k in p.get("answers", {}).keys())
    if not all_q_nums:
        return pd.DataFrame(), [], ["No answers found in extracted data"]
    max_q = max(all_q_nums)

    rows = []
    for p in participants:
        row = {
            "Name": str(p.get("name", "")),
            "Surname": str(p.get("surname", "")),
            "Language": str(p.get("language", "Unknown")),
        }
        answers = p.get("answers", {})
        for q in range(1, max_q + 1):
            row[f"Q{q}"] = str(answers.get(str(q), ""))
        rows.append(row)

    df = pd.DataFrame(rows)
    warnings = []
    errors = []

    if len(df) == 0:
        errors.append("No participants extracted")

    empty_names = (df["Name"].str.strip() == "").sum()
    if empty_names:
        warnings.append(f"{empty_names} participants with empty names")

    notes = data.get("notes", "")
    if notes:
        warnings.append(f"AI notes: {notes}")

    return df, warnings, errors


# ═══════════════════════════════════════════════════════════════
# API CALL
# ═══════════════════════════════════════════════════════════════

def call_api(prompt, api_key):
    """Call LLM via unified client."""
    from llm_client import call_llm
    provider = os.environ.get("CHIDON_PROVIDER", "anthropic")
    model = os.environ.get("CHIDON_MODEL", DEFAULT_MODEL)
    return call_llm(prompt, provider=provider, model=model, api_key=api_key)


def _parse_json_response(text):
    text = text.strip()
    text = re.sub(r'^```(?:json)?\s*', '', text)
    text = re.sub(r'\s*```$', '', text)
    return json.loads(text)


# ═══════════════════════════════════════════════════════════════
# ANSWER CLEANUP (normalize MC and T/F answers)
# ═══════════════════════════════════════════════════════════════

# T/F variants in multiple languages
_TRUE_VARIANTS = {"true", "t", "richtig", "r", "vrai", "v", "верно", "в",
                  "vero", "verdadero", "igaz", "pravda", "správně", "נכון"}
_FALSE_VARIANTS = {"false", "f", "falsch", "faux", "неверно", "н",
                   "falso", "hamis", "nepravda", "špatně", "לא נכון"}

# Pattern: letter optionally followed by punctuation and text
_MC_PATTERN = re.compile(r'^([A-Da-d])\s*[.):\-–—]\s*.*$')


def _clean_single_answer(val):
    """Normalize one answer value: T/F → T/F, MC → letter only."""
    val = str(val).strip()
    if not val:
        return ""

    low = val.lower().strip().rstrip(".")

    # T/F normalization
    if low in _TRUE_VARIANTS:
        return "T"
    if low in _FALSE_VARIANTS:
        return "F"

    # MC: "A. text" → "A", "b)" → "B"
    m = _MC_PATTERN.match(val)
    if m:
        return m.group(1).upper()

    # MC: single letter A-D
    if low in ("a", "b", "c", "d"):
        return val.upper()

    return val


def clean_answers_df(df):
    """Clean all Q* columns: normalize T/F and MC answers."""
    df = df.copy()
    q_cols = [c for c in df.columns if c.startswith("Q") and c[1:].isdigit()]
    for col in q_cols:
        df[col] = df[col].apply(_clean_single_answer)
    return df


# ═══════════════════════════════════════════════════════════════
# VALIDATION
# ═══════════════════════════════════════════════════════════════

def validate_parsed_answers(df):
    warnings, errors = [], []
    required = {"Name", "Surname", "Language"}
    missing = required - set(df.columns)
    if missing:
        errors.append(f"Missing columns: {missing}")
    q_cols = [c for c in df.columns if c.startswith("Q") and c[1:].isdigit()]
    if not q_cols:
        errors.append("No question columns")
    else:
        q_nums = sorted(int(c[1:]) for c in q_cols)
        gaps = set(range(q_nums[0], q_nums[-1] + 1)) - set(q_nums)
        if gaps:
            warnings.append(f"Missing questions: {sorted(gaps)}")
    if len(df) == 0:
        errors.append("No participants")
    return warnings, errors


def validate_parsed_key(df):
    warnings, errors = [], []
    required = {"Number", "Question", "Answer", "Points"}
    missing = required - set(df.columns)
    if missing:
        errors.append(f"Missing columns: {missing}")
    if len(df) == 0:
        errors.append("No questions")
    if "Answer" in df.columns:
        empty = (df["Answer"].str.strip() == "").sum()
        if empty:
            warnings.append(f"{empty} questions with empty answers")
    return warnings, errors


# ═══════════════════════════════════════════════════════════════
# PUBLIC API
# ═══════════════════════════════════════════════════════════════

def parse_answers_file(filepath, api_key):
    """
    Parse ANY file into clean Answers.xlsx format.
    Returns (df, mapping_or_info, warnings, errors).
    """
    ftype = get_file_type(filepath)

    if ftype == "tabular":
        # Tabular path: read → AI maps columns → deterministic transform
        df, metadata = read_tabular(filepath)
        headers, samples = extract_structure(df)
        prompt = _tabular_answers_prompt(headers, samples, metadata)
        response = call_api(prompt, api_key)
        mapping = _parse_json_response(response)
        clean_df = apply_answers_mapping(df, mapping)
        clean_df = clean_answers_df(clean_df)
        warnings, errors = validate_parsed_answers(clean_df)
        return clean_df, mapping, warnings, errors

    else:
        # Document path: extract text → AI returns structured data
        text, fmt = extract_text(filepath)
        if not text.strip():
            return pd.DataFrame(), {}, [], ["No text could be extracted from the file"]

        filename = os.path.basename(filepath)
        data = _json_from_document(text, filename, "answers", api_key)
        clean_df, warnings, errors = document_to_answers_df(data)
        clean_df = clean_answers_df(clean_df)
        info = {"source_format": fmt, "text_length": len(text),
                "extracted_participants": len(data.get("participants", [])),
                "ai_notes": data.get("notes", "")}
        return clean_df, info, warnings, errors


def parse_key_file(filepath, api_key):
    """
    Parse ANY file into clean Key.xlsx format.
    Returns (df, mapping_or_info, warnings, errors).
    """
    ftype = get_file_type(filepath)

    if ftype == "tabular":
        df, metadata = read_tabular(filepath)
        headers, samples = extract_structure(df)
        prompt = _tabular_key_prompt(headers, samples, metadata)
        response = call_api(prompt, api_key)
        mapping = _parse_json_response(response)
        clean_df = apply_key_mapping(df, mapping)
        warnings, errors = validate_parsed_key(clean_df)
        return clean_df, mapping, warnings, errors

    else:
        text, fmt = extract_text(filepath)
        if not text.strip():
            return pd.DataFrame(), {}, [], ["No text could be extracted from the file"]

        filename = os.path.basename(filepath)
        data = _json_from_document(text, filename, "key", api_key)
        clean_df, warnings, errors = document_to_key_df(data)
        info = {"source_format": fmt, "text_length": len(text),
                "extracted_questions": len(data.get("questions", [])),
                "ai_notes": data.get("notes", "")}
        return clean_df, info, warnings, errors


def detect_file_type(filepath, api_key):
    """Auto-detect whether file is a Key or Answers file."""
    ftype = get_file_type(filepath)

    if ftype == "tabular":
        df, metadata = read_tabular(filepath)
        headers, samples = extract_structure(df, max_rows=2)
        prompt = f"""What type of contest data is in this file?
FILE: "{metadata['filename']}" ({metadata['rows']} rows, {metadata['cols']} columns)
HEADERS: {json.dumps(headers, ensure_ascii=False)}
SAMPLE (first 2 rows): {json.dumps(samples[:2], ensure_ascii=False, indent=2)}

Reply ONLY with JSON: {{"type": "key" or "answers", "confidence": "high"/"medium"/"low", "reason": "..."}}"""
    else:
        text, _ = extract_text(filepath)
        snippet = text[:3000]
        prompt = f"""What type of contest data is in this document?
First 3000 chars:
---
{snippet}
---

Reply ONLY with JSON: {{"type": "key" or "answers", "confidence": "high"/"medium"/"low", "reason": "..."}}"""

    response = call_api(prompt, api_key)
    result = _parse_json_response(response)
    return result["type"], result["confidence"], result.get("reason", "")


# ═══════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════

def main():
    if len(sys.argv) < 3:
        print("Usage: python3 stage0.py <mode> <input_file> [output_file]")
        print("  mode: 'answers', 'key', or 'auto'")
        print(f"  formats: {', '.join(sorted(TABULAR_EXTENSIONS | DOCUMENT_EXTENSIONS))}")
        sys.exit(1)

    mode = sys.argv[1].lower()
    input_file = sys.argv[2]
    output_file = sys.argv[3] if len(sys.argv) > 3 else None

    api_key = os.environ.get("CHIDON_API_KEY", os.environ.get("ANTHROPIC_API_KEY", ""))
    if not api_key:
        print("ERROR: Set CHIDON_API_KEY or ANTHROPIC_API_KEY")
        sys.exit(1)

    if not os.path.exists(input_file):
        print(f"ERROR: File not found: {input_file}")
        sys.exit(1)

    ext = os.path.splitext(input_file)[1].lower()
    print(f"File: {input_file} ({ext})")
    print(f"Path: {'tabular' if ext in TABULAR_EXTENSIONS else 'document'}")

    if mode == "auto":
        print("Detecting file type...")
        file_type, confidence, reason = detect_file_type(input_file, api_key)
        print(f"  → {file_type} (confidence: {confidence}): {reason}")
        mode = file_type

    if mode == "answers":
        out = output_file or "Answers.xlsx"
        print(f"Parsing answers...")
        df, info, warnings, errors = parse_answers_file(input_file, api_key)
    elif mode == "key":
        out = output_file or "Key.xlsx"
        print(f"Parsing key...")
        df, info, warnings, errors = parse_key_file(input_file, api_key)
    else:
        print(f"Unknown mode: {mode}")
        sys.exit(1)

    print(f"\nParsing info:")
    print(json.dumps(info, ensure_ascii=False, indent=2))

    for w in warnings:
        print(f"  ⚠ {w}")
    for e in errors:
        print(f"  ❌ {e}")

    if errors:
        print("\nParsing failed.")
        sys.exit(1)

    print(f"\nResult: {len(df)} rows, {len(df.columns)} columns")
    print(df.head())

    df.to_excel(out, index=False)
    print(f"\nSaved: {out}")


if __name__ == "__main__":
    main()
