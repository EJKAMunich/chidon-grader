#!/usr/bin/env python3
"""
ЭТАП 5: Генерация Word-документов
Вход:  stage4_totals.json + Key.xlsx
Выход: ranking.docx + report_cards.docx

Формат карточки (по оригинальному шаблону):
  Part 1: Компактно — 2 строки (T/F + MC), ошибки красным
  Part 2: Подробно — Q21-50, эмодзи-маркеры, перевод нелатинских ответов

Запуск:
  python3 stage5.py
"""

import json
import openpyxl
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ═══════════════════════════════════════════════════════════════
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# ═══════════════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex, qn('w:val'): 'clear'
    })
    shading.append(shading_elem)


def set_cell_text(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def clear_cell(cell, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Очищает ячейку и возвращает paragraph для ручного заполнения."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    return p


def load_full_key(filepath):
    wb = openpyxl.load_workbook(filepath)
    ws = wb[wb.sheetnames[0]]
    key = {}
    for row in range(2, ws.max_row + 1):
        number = ws.cell(row=row, column=1).value
        if number is None:
            continue
        number = int(number)
        answer = str(ws.cell(row=row, column=3).value or "").strip()
        # Определяем тип
        if answer.upper() in ("T", "F"):
            q_type = "tf"
        elif answer.upper() in ("A", "B", "C", "D"):
            q_type = "mc"
        else:
            q_type = "open"
        key[number] = {
            "text": str(ws.cell(row=row, column=2).value or ""),
            "answer": answer,
            "points": int(ws.cell(row=row, column=4).value or 0),
            "comment": str(ws.cell(row=row, column=5).value or ""),
            "type": q_type,
        }
    return key


def is_latin(text):
    """Текст на латинице (English/German)?"""
    if not text:
        return True
    latin = sum(1 for ch in text if ch.isalpha() and ord(ch) < 0x0250)
    total = sum(1 for ch in text if ch.isalpha())
    return total == 0 or latin / total > 0.7


# ═══════════════════════════════════════════════════════════════
# RANKING (MODE A)
# ═══════════════════════════════════════════════════════════════

def generate_ranking(results, max_scores, output_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    title = doc.add_heading('Chidon HaTanach — Ranking', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    max_auto = max_scores["tf"] + max_scores["mc"]
    max_total = max_scores["total"]

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Заголовки: Pos | Name | Q1-20 (70) | Q21-50 (180) | Total Score
    headers = ["Pos", "Name", f"Q1-20 ({max_auto})",
               f"Q21-50 ({max_scores['open']})", "Total Score"]

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, size=10,
                      align=WD_ALIGN_PARAGRAPH.CENTER, color=(255, 255, 255))
        set_cell_shading(cell, "2C3E50")

    medal_colors = {1: "FFD700", 2: "E8E8E8", 3: "DEB887"}

    for r in results:
        row = table.add_row()
        c = row.cells
        set_cell_text(c[0], r["rank"], size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(c[1], r["full_name"], size=9)
        set_cell_text(c[2], r["auto_total"], size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(c[3], r["open_total"], size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(c[4], r["grand_total"], bold=True, size=10,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

        if r["rank"] in medal_colors:
            for cell in c:
                set_cell_shading(cell, medal_colors[r["rank"]])

    widths = [Cm(1.2), Cm(6.5), Cm(2.5), Cm(2.5), Cm(2.5)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    avg = sum(r["grand_total"] for r in results) / len(results)
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run(f"Participants: {len(results)} | Average: {avg:.1f}/{max_total}").font.size = Pt(9)

    doc.save(output_path)


# ═══════════════════════════════════════════════════════════════
# REPORT CARDS (MODE B)
# ═══════════════════════════════════════════════════════════════

def build_sequence_with_highlights(cell, participant_result, full_key, q_range, key_type):
    """
    Заполняет ячейку последовательностью ответов.
    Правильные — обычным шрифтом. Ошибки — красным жирным.
    Возвращает score.
    """
    p = clear_cell(cell, align=WD_ALIGN_PARAGRAPH.LEFT)
    auto_scores = participant_result.get("auto_scores", {})
    score = 0

    for idx, q_num in enumerate(q_range):
        q_str = str(q_num)
        q_info = full_key[q_num]

        sd = auto_scores.get(q_str, {})
        p_answer = sd.get("participant_answer", "")
        is_correct = sd.get("correct", False)
        earned = sd.get("earned", 0)
        score += earned

        if key_type == "tf":
            display = p_answer.upper() if p_answer else "-"
        else:
            display = p_answer[0].upper() if p_answer else "-"

        # Префикс: "1." или "6."
        prefix = f"{q_num}."

        if is_correct:
            run = p.add_run(f"{prefix}{display}")
            run.font.size = Pt(7)
        else:
            # Номер обычным
            run = p.add_run(prefix)
            run.font.size = Pt(7)
            # Ошибка красным жирным
            run = p.add_run(f"*{display}*")
            run.font.size = Pt(7)
            run.font.color.rgb = RGBColor(178, 34, 34)
            run.bold = True

        # Разделитель
        if idx < len(q_range) - 1:
            run = p.add_run(", ")
            run.font.size = Pt(7)

    return score


def generate_report_cards(results, max_scores, full_key, output_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(9)

    max_total = max_scores["total"]
    max_tf = max_scores["tf"]
    max_mc = max_scores["mc"]
    max_open = max_scores["open"]

    # Диапазоны по типу
    tf_range = sorted([q for q, info in full_key.items() if info["type"] == "tf"])
    mc_range = sorted([q for q, info in full_key.items() if info["type"] == "mc"])
    open_range = sorted([q for q, info in full_key.items() if info["type"] == "open"])

    # Правильные последовательности (одинаковые для всех)
    tf_correct_str = ", ".join(f"{q}.{full_key[q]['answer'].upper()}" for q in tf_range)
    mc_correct_str = ", ".join(f"{q}.{full_key[q]['answer'].upper()}" for q in mc_range)

    for p_idx, r in enumerate(results):
        if p_idx > 0:
            doc.add_page_break()

        # ── Заголовок: 👤 Participant Report ──
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run("👤 Participant Report")
        run.font.size = Pt(14)
        run.bold = True

        # ── Метаданные ──
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.paragraph_format.space_after = Pt(6)

        meta.add_run("Name: ").font.size = Pt(10)
        run = meta.add_run(r["full_name"])
        run.font.size = Pt(10)
        run.bold = True

        meta.add_run(" | Language: ").font.size = Pt(10)
        meta.add_run(r["language"]).font.size = Pt(10)

        meta.add_run(" | Total Score: ").font.size = Pt(10)
        run = meta.add_run(f"{r['grand_total']} / {max_total}")
        run.font.size = Pt(10)
        run.bold = True

        # ════════════════════════════════════════
        # PART 1: Quick Answers (Q1–20)
        # ════════════════════════════════════════

        doc.add_heading('Part 1: Quick Answers (Q1–20)', level=3)

        t1 = doc.add_table(rows=1, cols=5)
        t1.style = 'Table Grid'

        for i, h in enumerate(["Range", "Type", "Participant's Sequence",
                                "Correct Sequence", "Score"]):
            cell = t1.rows[0].cells[i]
            set_cell_text(cell, h, bold=True, size=8,
                          align=WD_ALIGN_PARAGRAPH.CENTER, color=(255, 255, 255))
            set_cell_shading(cell, "2C3E50")

        # ── Строка T/F ──
        if tf_range:
            row = t1.add_row()
            c = row.cells
            set_cell_text(c[0], f"{tf_range[0]}–{tf_range[-1]}", size=8,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(c[1], "True/False", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

            tf_score = build_sequence_with_highlights(c[2], r, full_key, tf_range, "tf")

            set_cell_text(c[3], tf_correct_str, size=7)
            set_cell_text(c[4], f"{tf_score}/{max_tf}", size=9, bold=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER)

        # ── Строка MC ──
        if mc_range:
            row = t1.add_row()
            c = row.cells
            set_cell_text(c[0], f"{mc_range[0]}–{mc_range[-1]}", size=8,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(c[1], "Multi. Choice", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

            mc_score = build_sequence_with_highlights(c[2], r, full_key, mc_range, "mc")

            set_cell_text(c[3], mc_correct_str, size=7)
            set_cell_text(c[4], f"{mc_score}/{max_mc}", size=9, bold=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER)

        # Ширина Part 1
        w1 = [Cm(1.5), Cm(2.2), Cm(5.5), Cm(5.5), Cm(1.5)]
        for row in t1.rows:
            for i, w in enumerate(w1):
                row.cells[i].width = w

        # ════════════════════════════════════════
        # PART 2: Capital / Open Questions (Q21–50)
        # ════════════════════════════════════════

        doc.add_heading('Part 2: Capital / Open Questions (Q21–50)', level=3)

        t2 = doc.add_table(rows=1, cols=6)
        t2.style = 'Table Grid'

        for i, h in enumerate(["#", "Question Text", "Participant Answer",
                                "Correct Answer", "Points", "Comment"]):
            cell = t2.rows[0].cells[i]
            set_cell_text(cell, h, bold=True, size=8,
                          align=WD_ALIGN_PARAGRAPH.CENTER, color=(255, 255, 255))
            set_cell_shading(cell, "2C3E50")

        for q_num in open_range:
            q_info = full_key[q_num]
            q_str = str(q_num)

            row = t2.add_row()
            c = row.cells

            # # — номер вопроса
            set_cell_text(c[0], q_num, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

            # Question Text — формулировка из ключа
            q_text = q_info["text"]
            if len(q_text) > 90:
                q_text = q_text[:87] + "..."
            set_cell_text(c[1], q_text, size=7)

            # Participant Answer — с переводом для нелатинских
            participant_answer = r.get("open_answers", {}).get(q_str, "")
            open_score_data = r.get("open_scores", {}).get(q_str, {})

            if not participant_answer:
                display_answer = "—"
            else:
                display_answer = participant_answer
                # Перевод: если ответ нелатинский → добавляем перевод из AI
                if not is_latin(participant_answer):
                    translation = open_score_data.get("translation", "—")
                    if translation and translation != "—":
                        display_answer = f"{participant_answer} ({translation})"

            if len(display_answer) > 65:
                display_answer = display_answer[:62] + "..."
            set_cell_text(c[2], display_answer, size=7)

            # Correct Answer
            correct = q_info["answer"]
            if len(correct) > 40:
                correct = correct[:37] + "..."
            set_cell_text(c[3], correct, size=7)

            # Points — эмодзи-маркеры: ✅ 6 pts, ⚠️ 3 pts, ❌ 0 pts
            earned = open_score_data.get("points", 0)
            max_pts = q_info["points"]

            if earned == max_pts:
                marker = "✅"
                pts_color = (34, 139, 34)
            elif earned > 0:
                marker = "⚠️"
                pts_color = (218, 165, 32)
            else:
                marker = "❌"
                pts_color = (178, 34, 34)

            set_cell_text(c[4], f"{marker} {earned} pts", size=8,
                          align=WD_ALIGN_PARAGRAPH.CENTER, color=pts_color)

            # Comment
            comment = open_score_data.get("comment", "")
            if len(comment) > 35:
                comment = comment[:32] + "..."
            set_cell_text(c[5], comment, size=7)

        # Ширина Part 2
        w2 = [Cm(0.8), Cm(4.5), Cm(3.8), Cm(3), Cm(1.8), Cm(3)]
        for row in t2.rows:
            for i, w in enumerate(w2):
                row.cells[i].width = w

        # ── Итоговая строка ──
        doc.add_paragraph("")
        summary = doc.add_paragraph()
        summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = summary.add_run(
            f"T/F: {r.get('tf_total', 0)}/{max_tf} | "
            f"MC: {r.get('mc_total', 0)}/{max_mc} | "
            f"Open: {r['open_total']}/{max_open} | "
            f"TOTAL: {r['grand_total']}/{max_total}"
        )
        run.bold = True
        run.font.size = Pt(11)

    doc.save(output_path)


# ═══════════════════════════════════════════════════════════════
# ЗАПУСК
# ═══════════════════════════════════════════════════════════════

def main():
    totals_file = "stage4_totals.json"
    key_file = "Key.xlsx"
    ranking_output = "ranking.docx"
    cards_output = "report_cards.docx"

    print("Читаю данные...")
    with open(totals_file, "r", encoding="utf-8") as f:
        data = json.load(f)

    results = data["results"]
    max_scores = data["max_scores"]
    full_key = load_full_key(key_file)

    print(f"  Участников: {len(results)}")
    print(f"  Вопросов: {len(full_key)}")

    print("Генерирую рейтинг...")
    generate_ranking(results, max_scores, ranking_output)
    print(f"  ✓ {ranking_output}")

    print("Генерирую карточки участников...")
    generate_report_cards(results, max_scores, full_key, cards_output)
    print(f"  ✓ {cards_output}")

    print(f"\nГотово!")


if __name__ == "__main__":
    main()
